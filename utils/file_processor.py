import os
import logging
import io
from werkzeug.utils import secure_filename

# Set up logging
logger = logging.getLogger(__name__)

# Try to import PDF and DOCX libraries, but don't fail if they're not available
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    logger.warning("PyPDF2 not installed. PDF text extraction will be limited.")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    logger.warning("python-docx not installed. DOCX text extraction will be limited.")

def extract_text_from_file(file_path):
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text or error message
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return f"[Error: File not found: {file_path}]"
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        # Plain text files
        if file_extension == '.txt':
            return extract_text_from_txt(file_path)
        
        # PDF files
        elif file_extension == '.pdf':
            if PDF_SUPPORT:
                return extract_text_from_pdf(file_path)
            else:
                return f"[PDF text extraction not available. Install PyPDF2 for full support: {os.path.basename(file_path)}]"
        
        # DOCX files
        elif file_extension == '.docx':
            if DOCX_SUPPORT:
                return extract_text_from_docx(file_path)
            else:
                return f"[DOCX text extraction not available. Install python-docx for full support: {os.path.basename(file_path)}]"
        
        # DOC files (old Word format)
        elif file_extension == '.doc':
            return f"[DOC format (pre-2007 Word) not directly supported. Please convert to DOCX: {os.path.basename(file_path)}]"
        
        # Unsupported file types
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return f"[Unsupported file type: {file_extension}]"
    
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}")
        return f"[Error extracting text: {str(e)}]"

def extract_text_from_txt(file_path):
    """Extract text from a plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file with latin-1 encoding: {str(e)}")
            return f"[Error reading text file: {str(e)}]"
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        return f"[Error reading text file: {str(e)}]"

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using PyPDF2"""
    if not PDF_SUPPORT:
        return f"[PDF text extraction not available. Install PyPDF2 for full support: {os.path.basename(file_path)}]"
    
    try:
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
        
        return "\n\n".join(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return f"[Error extracting text from PDF: {str(e)}]"

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx"""
    if not DOCX_SUPPORT:
        return f"[DOCX text extraction not available. Install python-docx for full support: {os.path.basename(file_path)}]"
    
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return f"[Error extracting text from DOCX: {str(e)}]"

def extract_text_from_folder(folder_path, max_chars=None):
    """
    Extract text from all supported files in a folder
    
    Args:
        folder_path (str): Path to the folder containing files
        max_chars (int, optional): Maximum number of characters to extract
        
    Returns:
        str: Concatenated text from all files
    """
    if not os.path.exists(folder_path):
        logger.warning(f"Folder not found: {folder_path}")
        return ""
    
    all_text = []
    total_chars = 0
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        # Skip directories
        if os.path.isdir(file_path):
            continue
        
        # Extract text from the file
        file_text = extract_text_from_file(file_path)
        file_size = len(file_text)
        
        logger.info(f"Extracted {file_size} characters from {filename}")
        
        # Add a header with the filename
        formatted_text = f"--- From {filename} ---\n{file_text}"
        all_text.append(formatted_text)
        
        total_chars += len(formatted_text)
        
        # Check if we've exceeded the maximum size
        if max_chars and total_chars > max_chars:
            logger.warning(f"Truncating extracted text at {max_chars} characters (current: {total_chars})")
            break
    
    combined_text = "\n\n".join(all_text)
    
    # Truncate if necessary
    if max_chars and len(combined_text) > max_chars:
        logger.warning(f"Truncating final combined text from {len(combined_text)} to {max_chars} characters")
        combined_text = combined_text[:max_chars] + "\n\n[Text truncated due to size limits]"
    
    logger.info(f"Total extracted text size: {len(combined_text)} characters from {len(all_text)} files")
    return combined_text

def get_file_size(file_path):
    """Get the size of a file in bytes"""
    try:
        return os.path.getsize(file_path)
    except Exception as e:
        logger.error(f"Error getting file size: {str(e)}")
        return 0

def get_folder_size(folder_path):
    """Get the total size of all files in a folder in bytes"""
    total_size = 0
    
    if not os.path.exists(folder_path):
        return 0
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        
        if os.path.isfile(file_path):
            total_size += get_file_size(file_path)
        elif os.path.isdir(file_path):
            total_size += get_folder_size(file_path)
    
    return total_size

def create_docx_from_text(text):
    """
    Create a DOCX file from text content
    
    Args:
        text (str): Text content to convert to DOCX
        
    Returns:
        BytesIO: BytesIO object containing the DOCX file
    """
    if not DOCX_SUPPORT:
        logger.error("python-docx not installed. Cannot create DOCX file.")
        return None
    
    try:
        # Create a new Document
        doc = docx.Document()
        
        # Split text by paragraphs and add each paragraph to the document
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():  # Skip empty paragraphs
                doc.add_paragraph(para)
        
        # Save the document to a BytesIO object
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes
    except Exception as e:
        logger.error(f"Error creating DOCX file: {str(e)}")
        return None
